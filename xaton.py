import json
from docx import Document
from docx.shared import Pt
import matplotlib.pyplot as plt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib.patches import Rectangle


# Функция для расчета центра массы грузов
def calculate_center_of_mass(data):
    total_mass = 0.0
    weighted_sum_x = 0.0
    weighted_sum_y = 0.0
    weighted_sum_z = 0.0

    for cargo in data.values():
        mass = cargo["вес"]
        total_mass += mass

        weighted_sum_x += mass * (cargo["длина"] / 2)
        weighted_sum_y += mass * (cargo["ширина"] / 2)
        weighted_sum_z += mass * (cargo["высота"] / 2)

    global center_of_mass_x
    center_of_mass_x = weighted_sum_x / total_mass
    global center_of_mass_y
    center_of_mass_y = weighted_sum_y / total_mass
    global center_of_mass_z
    center_of_mass_z = weighted_sum_z / total_mass

    return center_of_mass_x, center_of_mass_y, center_of_mass_z


# Функция для создания графика размещения грузов
def draw_cargo_placement(cargo_data, best_option):
    plt.figure(figsize=(8, 6))

    platform_width = 13.4
    platform_height = 2.72

    cargo_positions = []

    # Задаем начальное смещение
    x_offset = 0.0
    y_offset = 0.0

    cmap = plt.get_cmap('Blues')

    cargo_number = 1  # Начальный номер груза

    for cargo_name, cargo_info in cargo_data.items():
        cargo_length = cargo_info["длина"]
        cargo_width = cargo_info["ширина"]

        # Проверяем, помещается ли груз на текущем уровне платформы
        if x_offset + cargo_length <= platform_width and y_offset + cargo_width <= platform_height:
            cargo_x = x_offset
            cargo_y = y_offset
        else:
            # Если не помещается, переходим на следующий уровень платформы
            x_offset = 0.0
            y_offset += cargo_width
            cargo_x = x_offset
            cargo_y = y_offset

        plt.fill([cargo_x, cargo_x + cargo_length, cargo_x + cargo_length, cargo_x],
                 [cargo_y, cargo_y, cargo_y + cargo_width, cargo_y + cargo_width],
                 label=cargo_name)
        # Рисуем прямоугольник для груза с белым контуром
        rect = Rectangle((cargo_x, cargo_y), cargo_length, cargo_width, linewidth=2.5, edgecolor='white', facecolor=cmap(0.8))
        plt.gca().add_patch(rect)

        # Добавляем текстовую метку с номером груза
        plt.text(cargo_x + cargo_length / 2, cargo_y + cargo_width / 3, str(cargo_number), ha='center', va='center',
                 fontsize=15, color = 'white')

        # Обновляем смещение для следующего груза
        x_offset += cargo_length

        cargo_positions.append({
            "груз": cargo_name,
            "координаты": {"x": cargo_x, "y": cargo_y}
        })

        cargo_number += 1  # Увеличиваем номер груза

    known_point_x = center_of_mass_x
    known_point_y = center_of_mass_y
    plt.scatter(known_point_x, known_point_y, color='red', marker='o', label='Известная точка')
    # plt.legend()
    plt.title(f"Размещение грузов на платформе")
    plt.xlabel("Длина, м")
    plt.ylabel("Ширина, м")
    plt.grid(True)
    plt.axis([0, platform_width, 0, platform_height])
    plt.savefig("график_размещения.png")

    return cargo_positions

# Функция для создания документа Word
def create_word_document(data, center_of_mass, best_option, cargo_positions):
    doc = Document()
    doc.add_heading('Отчет по размещению грузов', 0)

    doc.add_heading('Информация о грузах', level=1)
    for cargo_name, cargo_info in data.items():
        doc.add_paragraph(f'{cargo_name}:')
        doc.add_paragraph(f'- Вес: {cargo_info["вес"]} т')
        doc.add_paragraph(f'- Длина: {cargo_info["длина"]} м')
        doc.add_paragraph(f'- Ширина: {cargo_info["ширина"]} м')
        doc.add_paragraph(f'- Высота: {cargo_info["высота"]} м')

    doc.add_heading('Результаты расчета', level=1)
    doc.add_paragraph(
        f'Центр массы грузов: ({center_of_mass[0]:.2f} м, {center_of_mass[1]:.2f} м, {center_of_mass[2]:.2f} м)')
    doc.add_paragraph(f'Наилучший вариант размещения: {best_option}')

    doc.add_heading('Размещение грузов', level=1)
    for cargo_position in cargo_positions:
        doc.add_paragraph(f'Груз: {cargo_position["груз"]}')
        doc.add_paragraph(
            f'Координаты (x, y): ({cargo_position["координаты"]["x"]:.2f} м, {cargo_position["координаты"]["y"]:.2f} м)')
    # Применяем стили ГОСТ Р 2.105-2019
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('отчет_размещение_грузов.docx')


if __name__ == "__main__":
    with open("данные.json", "r") as json_file:
        data = json.load(json_file)

    center_of_mass = calculate_center_of_mass(data)
    best_option = "на графике"  # Реализуйте логику выбора

    cargo_positions = draw_cargo_placement(data, best_option)
    create_word_document(data, center_of_mass, best_option, cargo_positions)

    print("Отчет успешно создан.")
